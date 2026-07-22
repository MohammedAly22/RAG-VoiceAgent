"""Ingestion pipeline: load → (multimodal) extract → chunk → embed → FAISS.

Supported formats: .txt .md .pdf .docx

For PDFs we use PyMuPDF (fitz): extract text per page, render a page screenshot
(shown in the UI when a chunk from that page is retrieved), pull embedded images
and caption them, and pull tables via pdfplumber. DOCX uses python-docx.

Run from the CLI:
    python -m backend.rag.ingest data/kb            # a folder
    python -m backend.rag.ingest file1.pdf file2.docx
"""
from __future__ import annotations

import logging
import re
import sys
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from config import CFG, abspath  # noqa: E402
from rag.store import get_store  # noqa: E402
from rag.multimodal import caption_image, read_page_image  # noqa: E402

logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s | INGEST | %(levelname)s | %(message)s",
                    datefmt="%H:%M:%S")
log = logging.getLogger("rag.ingest")

SUPPORTED = {".txt", ".md", ".pdf", ".docx"}
PAGES_DIR = Path(abspath(CFG.rag.pages_dir))


# ---------------------------------------------------------------------------
# chunking
# ---------------------------------------------------------------------------
def chunk_text(text: str, size: int, overlap: int) -> list[str]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    # split on paragraph boundaries first, then pack into ~size windows
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks, cur = [], ""
    for p in paras:
        if len(cur) + len(p) + 2 <= size:
            cur = f"{cur}\n\n{p}" if cur else p
        else:
            if cur:
                chunks.append(cur)
            if len(p) <= size:
                cur = p
            else:  # a single huge paragraph → hard window it
                for i in range(0, len(p), size - overlap):
                    chunks.append(p[i:i + size])
                cur = ""
    if cur:
        chunks.append(cur)
    # add overlap tails between consecutive chunks for context continuity
    if overlap > 0 and len(chunks) > 1:
        out = [chunks[0]]
        for prev, nxt in zip(chunks, chunks[1:]):
            tail = prev[-overlap:]
            out.append((tail + "\n" + nxt) if tail.strip() else nxt)
        chunks = out
    return chunks


def normalize_ar(text: str) -> str:
    """Normalize Arabic text extracted from PDFs.

    PDFs that embed *shaped* glyphs extract as Unicode presentation forms
    (ﻃﺒﻖ instead of طبق), which would never match a user's normally-typed query.
    NFKC folds those back to the base letters. We also strip tatweel and
    normalize a few common variants so retrieval is stable.
    """
    import re
    import unicodedata
    t = unicodedata.normalize("NFKC", text or "")
    t = t.replace("ـ", "")                       # tatweel
    t = re.sub(r"[‎‏‪-‮]", "", t)  # bidi control marks
    return t


_PRESENTATION = re.compile(r"[ﭐ-﷿ﹰ-﻿]")


def _text_layer_is_unreliable(text: str) -> bool:
    """True when a PDF's text layer is shaped/bidi-mangled (presentation forms).

    Such layers also tend to reverse digit runs (١٨٠ → ٠٨١), so prices/numbers
    can't be trusted — we prefer reading the rendered page with vision instead.
    """
    if not text:
        return False
    hits = len(_PRESENTATION.findall(text))
    return hits >= 8 or (hits and hits / max(len(text), 1) > 0.02)


def _mk(text: str, source: str, page, ctype: str, page_image, idx: int) -> dict:
    return {
        "id": uuid.uuid4().hex[:12],
        "text": text.strip(),
        "source": source,
        "page": page,
        "chunk_index": idx,
        "type": ctype,
        "page_image": page_image,
    }


# ---------------------------------------------------------------------------
# loaders → list of "raw units" (text/table/image) with page info
# ---------------------------------------------------------------------------
def load_pdf(path: Path) -> list[dict]:
    import fitz  # PyMuPDF
    src = path.name
    out: list[dict] = []
    page_dir = PAGES_DIR / src
    page_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(path)
    vision_pages: set[int] = set()   # pages whose text we recovered with vision
    for pno in range(len(doc)):
        page = doc[pno]
        pageno = pno + 1
        # render page screenshot for the UI
        rel_img = f"{src}/p{pageno}.png"
        pix = page.get_pixmap(dpi=120)
        pix.save(str(PAGES_DIR / rel_img))
        # text — but if the layer is shaped/bidi-mangled, read the rendered page
        # with vision instead (accurate Arabic + correct numbers).
        text = page.get_text("text").strip()
        if _text_layer_is_unreliable(text) and CFG.rag.multimodal:
            log.info("  page %d: text layer unreliable → reading with vision", pageno)
            vis = read_page_image(str(PAGES_DIR / rel_img))
            if vis:
                out.append({"text": vis, "page": pageno, "type": "text", "page_image": rel_img})
                vision_pages.add(pageno)
                text = ""   # don't also add the mangled layer
            else:
                text = normalize_ar(text)
        else:
            text = normalize_ar(text)
        if text:
            out.append({"text": text, "page": pageno, "type": "text", "page_image": rel_img})
        # embedded images → caption (multimodal)
        if CFG.rag.multimodal:
            for ii, img in enumerate(page.get_images(full=True)):
                try:
                    xref = img[0]
                    pm = fitz.Pixmap(doc, xref)
                    if pm.n >= 5:  # CMYK/alpha → RGB
                        pm = fitz.Pixmap(fitz.csRGB, pm)
                    tmp = str(PAGES_DIR / f"{src}/p{pageno}_img{ii}.png")
                    pm.save(tmp)
                    cap = caption_image(tmp)
                    if cap:
                        out.append({"text": f"[صورة صفحة {pageno}] {cap}",
                                    "page": pageno, "type": "image_caption",
                                    "page_image": rel_img})
                except Exception as e:  # noqa: BLE001
                    log.warning("image extract failed p%d: %s", pageno, e)
    doc.close()
    # tables via pdfplumber (better table structure than fitz). Skipped for pages
    # we already read with vision — their glyph-level text is mangled, and the
    # vision pass already captured those tables correctly.
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for pno, page in enumerate(pdf.pages):
                if (pno + 1) in vision_pages:
                    continue
                for tbl in page.extract_tables() or []:
                    rows = [" | ".join(normalize_ar((c or "").strip()) for c in row) for row in tbl if row]
                    if len(rows) >= 2:
                        out.append({"text": f"[جدول صفحة {pno+1}]\n" + "\n".join(rows),
                                    "page": pno + 1, "type": "table",
                                    "page_image": f"{src}/p{pno+1}.png"})
    except Exception as e:  # noqa: BLE001
        log.warning("table extraction failed: %s", e)
    return out


def load_docx(path: Path) -> list[dict]:
    import docx
    d = docx.Document(str(path))
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    out = [{"text": "\n".join(paras), "page": None, "type": "text", "page_image": None}]
    # tables
    for ti, tbl in enumerate(d.tables):
        rows = [" | ".join(c.text.strip() for c in row.cells) for row in tbl.rows]
        if len(rows) >= 2:
            out.append({"text": f"[جدول {ti+1}]\n" + "\n".join(rows),
                        "page": None, "type": "table", "page_image": None})
    return out


def load_text(path: Path) -> list[dict]:
    return [{"text": path.read_text(encoding="utf-8", errors="ignore"),
             "page": None, "type": "text", "page_image": None}]


def load_file(path: Path) -> list[dict]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return load_pdf(path)
    if ext == ".docx":
        return load_docx(path)
    if ext in {".txt", ".md"}:
        return load_text(path)
    raise ValueError(f"unsupported format: {ext}")


# ---------------------------------------------------------------------------
# top-level ingest
# ---------------------------------------------------------------------------
def ingest_file(path: str | Path) -> int:
    path = Path(path)
    if path.suffix.lower() not in SUPPORTED:
        log.warning("skip unsupported %s", path.name)
        return 0
    log.info("ingesting %s", path.name)
    store = get_store()
    store.remove_source(path.name)  # re-ingest replaces old chunks for this file
    units = load_file(path)
    chunks: list[dict] = []
    idx = 0
    for u in units:
        pieces = chunk_text(u["text"], CFG.rag.chunk_size, CFG.rag.chunk_overlap) \
            if u["type"] == "text" else [u["text"]]
        for piece in pieces:
            if piece.strip():
                chunks.append(_mk(piece, path.name, u["page"], u["type"],
                                  u["page_image"], idx))
                idx += 1
    n = store.add(chunks)
    log.info("  → %d chunks from %s", n, path.name)
    return n


def ingest_paths(paths: list[str]) -> int:
    total = 0
    for p in paths:
        p = Path(p)
        if p.is_dir():
            for f in sorted(p.iterdir()):
                if f.suffix.lower() in SUPPORTED:
                    total += ingest_file(f)
        elif p.exists():
            total += ingest_file(p)
        else:
            log.warning("not found: %s", p)
    return total


if __name__ == "__main__":
    args = sys.argv[1:] or ["data/kb"]
    total = ingest_paths([abspath(a) if not Path(a).is_absolute() else a for a in args])
    store = get_store()
    log.info("DONE. total chunks now in store: %d", store.count())
    print("\nSources:")
    for s in store.sources():
        print(f"  - {s['source']}: {s['chunks']} chunks, {s['pages']} pages, types={s['types']}")
